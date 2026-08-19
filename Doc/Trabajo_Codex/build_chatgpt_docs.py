from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
import sys

OUT = Path(sys.argv[1]); OUT.mkdir(parents=True, exist_ok=True)
NAVY="#07182F"; BLUE="#0B3B70"; GOLD="#EBA03D"; CYAN="#00D4FF"; LIGHT="#F4F7FA"

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement("w:shd"); shd.set(qn("w:fill"), fill.replace("#","")); tcPr.append(shd)

def base_doc(title, subtitle):
    d=Document(); sec=d.sections[0]; sec.top_margin=Inches(.7); sec.bottom_margin=Inches(.7); sec.left_margin=Inches(.8); sec.right_margin=Inches(.8)
    styles=d.styles
    normal=styles["Normal"]; normal.font.name="Calibri"; normal.font.size=Pt(10.5); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.15
    for n,size,color in [("Heading 1",16,GOLD),("Heading 2",13,BLUE),("Heading 3",11.5,BLUE)]:
        s=styles[n]; s.font.name="Calibri"; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color.replace("#","")); s.paragraph_format.space_before=Pt(10); s.paragraph_format.space_after=Pt(5)
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(title); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=RGBColor.from_string(GOLD.replace("#",""))
    p=d.add_paragraph(subtitle); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].italic=True; p.runs[0].font.color.rgb=RGBColor.from_string(BLUE.replace("#",""))
    d.add_paragraph("ZYRON AI LAB · Versión 2026-07-26").alignment=WD_ALIGN_PARAGRAPH.CENTER
    return d

def add_field_table(d, fields):
    t=d.add_table(rows=1, cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    t.columns[0].width=Inches(1.8); t.columns[1].width=Inches(4.7)
    h=t.rows[0].cells; h[0].merge(h[1]); h[0].text="REGISTRO EDITABLE"; shade(h[0], BLUE)
    for run in h[0].paragraphs[0].runs: run.font.color.rgb=RGBColor(255,255,255); run.bold=True
    for label, hint in fields:
        c=t.add_row().cells; c[0].text=label; c[1].text=hint
        shade(c[0], "E8EEF5"); c[0].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        c[0].paragraphs[0].runs[0].bold=True
        c[1].paragraphs[0].paragraph_format.space_after=Pt(9 if len(hint)<30 else 13)
    return t

prompt_fields=[
("Nombre del proyecto",""),("Fecha","AAAA-MM-DD"),("Responsable",""),("Objetivo","¿Qué resultado se necesita?"),
("Contexto","Antecedentes necesarios para comprender la tarea."),("Información disponible","Documentos, datos y referencias."),
("Tarea solicitada","Acción concreta que debe realizar ChatGPT."),("Restricciones","Límites, exclusiones, extensión, tono o condiciones."),
("Formato esperado","Estructura de la respuesta."),("Criterios de calidad","Condiciones para aprobar el resultado."),
("Fuentes disponibles","Fuentes que pueden consultarse."),("Datos que deben verificarse","Afirmaciones, cifras, fechas o fuentes."),
("Resultado inicial","Pegue o resuma la primera respuesta."),("Correcciones solicitadas","Cambios específicos pedidos."),
("Resultado final aprobado","Versión aceptada después de la revisión."),("Observaciones","")
]
d=base_doc("Plantilla de prompt estructurado","Formulario editable para planificar, revisar y aprobar una interacción con ChatGPT.")
add_field_table(d,prompt_fields); d.save(OUT/"plantilla-prompt-estructurado-chatgpt.docx")

security_sections=[
("1. Datos que no deben cargarse","No introducir contraseñas, secretos, credenciales, números completos de identificación, datos bancarios ni información cuya exposición pueda causar daño."),
("2. Protección de datos personales","Aplicar minimización, autorización, finalidad definida y controles acordes con la normativa y las políticas internas."),
("3. Información empresarial confidencial","No cargar contratos, estrategias, bases de clientes, datos internos o propiedad reservada sin autorización y controles aprobados."),
("4. Contraseñas y credenciales","Nunca compartir contraseñas, tokens, claves API, códigos de recuperación ni secretos de acceso."),
("5. Propiedad intelectual","Verificar permisos, autoría, licencias y condiciones antes de cargar o reutilizar contenido."),
("6. Verificación de respuestas","Contrastar hechos, cifras, fechas y citas con fuentes confiables, preferentemente oficiales."),
("7. Sesgos y alucinaciones","Una respuesta fluida puede ser incorrecta, incompleta o sesgada. Solicitar evidencia y revisión crítica."),
("8. Uso médico, legal y financiero","No utilizar la respuesta como sustituto de profesionales cualificados ni para decisiones de alto impacto sin supervisión."),
("9. Supervisión humana","Una persona responsable debe revisar y aprobar todo resultado institucional."),
("10. Antes de publicar","Confirmar exactitud, fuentes, privacidad, derechos, tono, vigencia y autorización."),
("11. Fuentes y versiones","Registrar URLs, documentos, fecha de consulta, versión del borrador y responsable de aprobación."),
("12. Respuesta incorrecta","Detener el uso, identificar el error, conservar evidencia, consultar fuentes, solicitar una corrección específica y someter la nueva versión a revisión.")
]
d=base_doc("Guía de seguridad y uso responsable","Referencia educativa para trabajar con ChatGPT de manera supervisada.")
d.sections[0].top_margin=Inches(.45); d.sections[0].bottom_margin=Inches(.45)
d.styles["Normal"].font.size=Pt(9); d.styles["Normal"].paragraph_format.space_after=Pt(2); d.styles["Normal"].paragraph_format.line_spacing=1
d.styles["Heading 1"].font.size=Pt(11); d.styles["Heading 1"].paragraph_format.space_before=Pt(5); d.styles["Heading 1"].paragraph_format.space_after=Pt(2)
for h,b in security_sections: d.add_heading(h,level=1); d.add_paragraph(b)
p=d.add_paragraph(); r=p.add_run("Nota: "); r.bold=True; p.add_run("Esta guía es educativa y no sustituye las políticas internas, la asesoría jurídica ni los protocolos de seguridad de cada organización.")
d.save(OUT/"guia-seguridad-chatgpt.docx")

styles=getSampleStyleSheet()
title_style=ParagraphStyle("ZTitle",parent=styles["Title"],fontName="Helvetica-Bold",fontSize=20,textColor=colors.HexColor(GOLD),alignment=TA_CENTER,spaceAfter=12)
h_style=ParagraphStyle("ZH",parent=styles["Heading2"],fontName="Helvetica-Bold",fontSize=12,textColor=colors.HexColor(BLUE),spaceBefore=8,spaceAfter=4)
body_style=ParagraphStyle("ZBody",parent=styles["BodyText"],fontName="Helvetica",fontSize=9.3,leading=12,textColor=colors.HexColor(NAVY))
small_style=ParagraphStyle("ZSmall",parent=body_style,fontSize=7.5,leading=9)
header_style=ParagraphStyle("ZHeader",parent=small_style,textColor=colors.white,alignment=TA_CENTER)

def pdf_header(title, subtitle):
    return [Paragraph(title,title_style),Paragraph(subtitle,ParagraphStyle("sub",parent=body_style,alignment=TA_CENTER,textColor=colors.HexColor(BLUE))),Spacer(1,10)]

def make_prompt_pdf():
    story=pdf_header("Plantilla de prompt estructurado","Laboratorio ChatGPT · ZYRON AI LAB · Versión 2026-07-26")
    data=[[Paragraph("<b>Campo</b>",header_style),Paragraph("<b>Contenido</b>",header_style)]]
    for a,b in prompt_fields:data.append([Paragraph(a,small_style),Paragraph((b+"<br/><br/>") if b else "<br/><br/>",small_style)])
    t=Table(data,colWidths=[1.65*72,4.85*72],repeatRows=1)
    t.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor(BLUE)),("TEXTCOLOR",(0,0),(-1,0),colors.white),("BACKGROUND",(0,1),(0,-1),colors.HexColor("#E8EEF5")),("GRID",(0,0),(-1,-1),.4,colors.HexColor("#AAB8C5")),("VALIGN",(0,0),(-1,-1),"TOP"),("LEFTPADDING",(0,0),(-1,-1),6),("RIGHTPADDING",(0,0),(-1,-1),6),("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5)]))
    story.append(t); SimpleDocTemplate(str(OUT/"plantilla-prompt-estructurado-chatgpt.pdf"),pagesize=letter,rightMargin=36,leftMargin=36,topMargin=36,bottomMargin=36).build(story)

def make_security_pdf():
    story=pdf_header("Guía de seguridad y uso responsable","Laboratorio ChatGPT · ZYRON AI LAB · Versión 2026-07-26")
    for h,b in security_sections: story += [Paragraph(h,h_style),Paragraph(b,body_style)]
    story += [Spacer(1,8),Paragraph("<b>Nota:</b> Esta guía es educativa y no sustituye las políticas internas, la asesoría jurídica ni los protocolos de seguridad de cada organización.",body_style)]
    SimpleDocTemplate(str(OUT/"guia-seguridad-chatgpt.pdf"),pagesize=letter,rightMargin=46,leftMargin=46,topMargin=40,bottomMargin=40).build(story)

def make_grid_pdf(filename,title,headers,rows,widths,orientation=landscape(letter)):
    story=pdf_header(title,"Laboratorio ChatGPT · ZYRON AI LAB · Versión 2026-07-26")
    data=[[Paragraph(f"<b>{x}</b>",header_style) for x in headers]]+[[Paragraph(str(x),small_style) for x in row] for row in rows]
    t=Table(data,colWidths=widths,repeatRows=1)
    t.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor(BLUE)),("TEXTCOLOR",(0,0),(-1,0),colors.white),("GRID",(0,0),(-1,-1),.35,colors.HexColor("#AAB8C5")),("VALIGN",(0,0),(-1,-1),"TOP"),("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white,colors.HexColor(LIGHT)]),("LEFTPADDING",(0,0),(-1,-1),4),("RIGHTPADDING",(0,0),(-1,-1),4),("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5)]))
    story.append(t); SimpleDocTemplate(str(OUT/filename),pagesize=orientation,rightMargin=24,leftMargin=24,topMargin=30,bottomMargin=30).build(story)

criteria=["Objetivo definido","Contexto suficiente","Restricciones claras","Formato solicitado","Protección de datos","Verificación de hechos","Revisión de fuentes","Información desactualizada identificada","Corrección de errores","Aprobación humana","Registro de versiones","Calidad final del entregable"]
make_grid_pdf("lista-verificacion-chatgpt.pdf","Lista de verificación para utilizar ChatGPT",["N°","Criterio","Cumple","No cumple","No aplica","Observaciones"],[[i+1,c,"","","",""] for i,c in enumerate(criteria)],[28,210,55,62,55,310])

rubrics=[("Definición del problema",15),("Calidad de las instrucciones",20),("Uso del contexto",15),("Revisión e iteración",15),("Verificación de información",15),("Seguridad y privacidad",10),("Calidad del entregable",10)]
make_grid_pdf("rubrica-laboratorio-chatgpt.pdf","Rúbrica de evaluación del laboratorio",["Criterio","Máx.","Excelente","Competente","En desarrollo","Requiere revisión","Obtenido"],[[a,b,"Dominio completo","Cumple lo esencial","Cumplimiento parcial","Debe rehacerse",""] for a,b in rubrics],[120,36,125,125,125,125,50])

iteration_headers=["N°","Fecha","Objetivo","Prompt utilizado","Resultado recibido","Error o vacío","Corrección","Fuente","Validado","Responsable","Estado","Observaciones"]
make_grid_pdf("registro-iteraciones-chatgpt.pdf","Registro de iteraciones con ChatGPT",iteration_headers,[["","","","","","","","","","","",""] for _ in range(8)],[25,48,70,95,90,80,80,75,45,65,55,70])
make_prompt_pdf(); make_security_pdf()
