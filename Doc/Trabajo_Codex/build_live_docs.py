from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak

ROOT=Path(r"C:\Users\ASUS\Desktop\Empresa Zyron\ZYRON PRIMUS")
OUT=ROOT/"recursos"/"laboratorio-live"
OUT.mkdir(parents=True,exist_ok=True)
NAVY="07152B"; BLUE="0B3F75"; GOLD="F0A22E"; LIGHT="EEF3F8"

def shade(cell, fill):
    p=cell._tc.get_or_add_tcPr(); x=OxmlElement("w:shd"); x.set(qn("w:fill"),fill); p.append(x)

def doc(title, subtitle):
    d=Document(); sec=d.sections[0]
    sec.page_height,sec.page_width=Inches(11),Inches(8.5)
    sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1)
    n=d.styles["Normal"]; n.font.name="Arial"; n.font.size=Pt(10.5); n.font.color.rgb=RGBColor.from_string(NAVY)
    n.paragraph_format.space_after=Pt(6); n.paragraph_format.line_spacing=1.15
    for k,size,color in [("Heading 1",16,BLUE),("Heading 2",13,BLUE),("Heading 3",11,NAVY)]:
        s=d.styles[k]; s.font.name="Arial"; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color)
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(title); r.bold=True; r.font.name="Arial"; r.font.size=Pt(23); r.font.color.rgb=RGBColor.from_string(GOLD)
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run(subtitle).italic=True
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run("ZYRON AI LAB · Revisión 26 de julio de 2026").bold=True
    return d

def table(d, headers, rows, widths=None):
    t=d.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    for i,h in enumerate(headers):
        t.rows[0].cells[i].text=h; shade(t.rows[0].cells[i],BLUE)
        for run in t.rows[0].cells[i].paragraphs[0].runs: run.bold=True; run.font.color.rgb=RGBColor(255,255,255)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row): cells[i].text=str(v); shade(cells[i], LIGHT if len(t.rows)%2==0 else "FFFFFF")
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    return t

steps=[
("1. Definir el objetivo","Escriba qué resultado necesita y cómo comprobará que la sesión fue útil."),
("2. Verificar disponibilidad","Confirme plan, región, aplicación, cuenta y opciones visibles en Configuración > Voice."),
("3. Elegir la opción","Use Live para diálogo natural, Advanced para funciones móviles compatibles o Standard para turnos."),
("4. Preparar el entorno","Busque un lugar silencioso, use auriculares cuando convenga y estabilice la conexión."),
("5. Revisar permisos","Autorice solo micrófono y, si corresponde, cámara o pantalla durante el tiempo necesario."),
("6. Proteger información","Retire datos confidenciales, notificaciones y ventanas ajenas; obtenga consentimiento de terceros."),
("7. Dar contexto","Indique propósito, audiencia, límites, idioma y formato de la respuesta."),
("8. Conversar por bloques","Hable con frases claras, pause y corrija de inmediato cualquier interpretación errónea."),
("9. Interrumpir con intención","Detenga o redirija la respuesta cuando se aparte del objetivo; el ruido puede provocar interrupciones."),
("10. Validar","Pida un resumen, contraste datos importantes con fuentes y no trate la transcripción como acta literal."),
("11. Cerrar","Solicite conclusiones, próximos pasos y asuntos sin resolver; detenga micrófono, cámara o pantalla."),
("12. Registrar","Revise historial y transcripción; documente resultado, incidencias, decisiones y eliminación necesaria.")
]
guide=doc("Guía de preparación para una sesión Voice/Live","Preparación, conducción, validación y cierre supervisado")
guide.add_heading("Propósito",1); guide.add_paragraph("Preparar una sesión responsable de ChatGPT Voice. Live es una opción dentro de Voice; su disponibilidad y funciones dependen del plan, la cuenta, la región, la aplicación y la configuración.")
guide.add_heading("Antes de comenzar",1)
for x in ["Objetivo y criterio de éxito definidos.","Dispositivo cargado, red estable, micrófono probado y salida de audio adecuada.","Permisos mínimos y consentimiento de terceros confirmados.","Información confidencial retirada o protegida.","Opción de Voice verificada en la cuenta real."]: guide.add_paragraph(x,style="List Bullet")
guide.add_heading("Tutorial de doce pasos",1); table(guide,["Paso","Acción"],steps,[1.6,4.7])
guide.add_heading("Durante la sesión",1)
for x in ["Mantenga supervisión humana y confirme fechas, nombres, cifras y decisiones.","Evite conversaciones con varias personas: Live está pensado principalmente para diálogo individual.","La transcripción puede diferir del audio por ruido, solapamiento o velocidad.","Voice puede coordinar Work o Codex en escritorio cuando esté disponible; Work y Codex ejecutan la tarea especializada."]: guide.add_paragraph(x,style="List Bullet")
guide.add_heading("Cierre seguro",1)
for x in ["Detenga inmediatamente la sesión si aparece información no autorizada.","Revise la conversación y corrija el registro antes de usarlo.","Elimine la conversación cuando sea necesario conforme a controles de datos y políticas internas.","No es obligatorio conservar audio ni video como evidencia del laboratorio."]: guide.add_paragraph(x,style="List Bullet")
guide.add_heading("Fuentes oficiales",1)
for u in ["https://help.openai.com/en/articles/20001274","https://help.openai.com/en/articles/8400625-voice-mode","https://help.openai.com/en/articles/12168547-voice-dictation-faq","https://help.openai.com/en/articles/20001275-chatgpt-work-and-codex"]: guide.add_paragraph(u)
guide.save(OUT/"guia-preparacion-voice-live.docx")

security=doc("Guía de seguridad para voz, cámara y pantalla","Privacidad, consentimiento y respuesta ante incidentes")
sections=[
("1. Micrófono","Autorice el micrófono solo en el dispositivo y cuenta correctos. Silencie o finalice al terminar."),
("2. Cámara y pantalla","No son funciones generales de Live. Cuando Advanced u otra experiencia elegible las ofrezca, comparta únicamente el área necesaria y cierre notificaciones."),
("3. Consentimiento","Informe a terceros antes de captar su voz, imagen o información. No continúe sin autorización válida."),
("4. Información hablada","Evite credenciales, secretos, datos personales innecesarios, información médica, financiera, contractual o restringida."),
("5. Espacios públicos","Use auriculares, reduzca el volumen y no revele información sensible donde terceros puedan oír."),
("6. Retención y controles","Revise los controles de datos vigentes. Audio, video y transcripciones pueden tener reglas distintas; eliminar un chat puede iniciar procesos de eliminación sujetos a excepciones."),
("7. Transcripciones","No son actas literales. Compare con notas humanas y fuentes antes de aprobar decisiones."),
("8. Segundo plano","Solo cuando la aplicación, configuración, límites y estado del dispositivo lo permitan. No prometa continuidad ilimitada."),
("9. Work y Codex","Voice coordina; los agentes especializados ejecutan. Revise permisos, progreso, bloqueos y entregables."),
("10. Incidente","Detenga la sesión, silencie permisos, conserve evidencia mínima, informe al responsable, corrija el registro y elimine datos cuando corresponda."),
("11. Responsabilidad","Una persona competente valida exactitud, privacidad, autorización y uso final.")
]
for h,b in sections: security.add_heading(h,1); security.add_paragraph(b)
security.add_heading("Lista de salida",1)
for x in ["Micrófono detenido","Cámara o pantalla detenida","Terceros informados","Transcripción revisada","Datos sensibles tratados","Resultado validado","Conversación eliminada si procede"]: security.add_paragraph(x,style="List Bullet")
security.add_paragraph("Material educativo: no sustituye políticas internas, asesoría jurídica ni controles profesionales.")
security.save(OUT/"guia-seguridad-voz-camara-pantalla.docx")

styles=getSampleStyleSheet()
T=ParagraphStyle("T",parent=styles["Title"],fontName="Helvetica-Bold",fontSize=20,textColor=colors.HexColor("#"+GOLD),alignment=1,spaceAfter=8)
H=ParagraphStyle("H",parent=styles["Heading1"],fontName="Helvetica-Bold",fontSize=11,textColor=colors.HexColor("#"+BLUE),spaceBefore=5,spaceAfter=3)
B=ParagraphStyle("B",parent=styles["BodyText"],fontName="Helvetica",fontSize=8.4,leading=10,textColor=colors.HexColor("#"+NAVY),spaceAfter=3)

def pdf_doc(path,title,blocks):
    story=[Paragraph(title,T),Paragraph("ZYRON AI LAB · Revisión 26 de julio de 2026",ParagraphStyle("S",parent=B,alignment=1,fontSize=8)),Spacer(1,6)]
    for h,items in blocks:
        story.append(Paragraph(h,H))
        if isinstance(items,str): story.append(Paragraph(items,B))
        else:
            for item in items: story.append(Paragraph("• "+item,B))
    SimpleDocTemplate(str(path),pagesize=letter,rightMargin=48,leftMargin=48,topMargin=40,bottomMargin=40).build(story)

pdf_doc(OUT/"guia-preparacion-voice-live.pdf","Guía de preparación para una sesión Voice/Live",[
("Propósito","Preparar, conducir, validar y cerrar una sesión supervisada. Live es una opción dentro de ChatGPT Voice y su disponibilidad es condicionada."),
("Antes de comenzar",["Objetivo y criterio de éxito definidos.","Dispositivo, red, micrófono y audio probados.","Permisos mínimos y consentimiento confirmados.","Información confidencial retirada.","Opción real de Voice verificada."]),
("Tutorial de doce pasos",[f"<b>{a}</b>: {b}" for a,b in steps]),
("Controles esenciales",["Valide datos y decisiones.","No trate la transcripción como acta literal.","Voice puede coordinar Work/Codex; los agentes ejecutan.","No es obligatorio guardar audio o video.","Detenga y elimine cuando sea necesario."]),
("Fuentes oficiales revisadas el 26-07-2026",["ChatGPT Voice: https://help.openai.com/en/articles/20001274","Voice Mode FAQ: https://help.openai.com/en/articles/8400625-voice-mode","Voice Dictation FAQ: https://help.openai.com/en/articles/12168547-voice-dictation-faq","ChatGPT Work and Codex: https://help.openai.com/en/articles/20001275-chatgpt-work-and-codex"])
])
pdf_doc(OUT/"guia-seguridad-voz-camara-pantalla.pdf","Guía de seguridad para voz, cámara y pantalla",[(h,b) for h,b in sections]+[("Lista de salida",["Micrófono detenido.","Cámara o pantalla detenida.","Consentimiento verificado.","Transcripción revisada.","Datos tratados.","Resultado validado.","Conversación eliminada si procede."]),("Alcance","Material educativo: no sustituye políticas internas, asesoría jurídica ni controles profesionales.")])

def grid_pdf(path,title,headers,rows,widths):
    st=[Paragraph(title,T),Paragraph("ZYRON AI LAB · 26 de julio de 2026",ParagraphStyle("SG",parent=B,alignment=1,fontSize=8)),Spacer(1,6)]
    hs=ParagraphStyle("TH",parent=B,textColor=colors.white,fontName="Helvetica-Bold",alignment=1,fontSize=7)
    bs=ParagraphStyle("TB",parent=B,fontSize=6.5,leading=8)
    data=[[Paragraph(str(x),hs) for x in headers]]+[[Paragraph(str(x),bs) for x in row] for row in rows]
    t=Table(data,colWidths=widths,repeatRows=1)
    t.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor("#"+BLUE)),("GRID",(0,0),(-1,-1),.35,colors.HexColor("#AAB8C5")),("VALIGN",(0,0),(-1,-1),"TOP"),("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white,colors.HexColor("#"+LIGHT)]),("LEFTPADDING",(0,0),(-1,-1),3),("RIGHTPADDING",(0,0),(-1,-1),3),("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3)]))
    st.append(t); SimpleDocTemplate(str(path),pagesize=letter,rightMargin=24,leftMargin=24,topMargin=28,bottomMargin=28).build(st)

checks=["Objetivo y criterio de éxito definidos","Plan, región, aplicación y cuenta verificados","Opción Voice/Live/Advanced/Standard confirmada","Conexión y batería adecuadas","Micrófono probado","Auriculares o salida de audio adecuada","Espacio silencioso y privado","Permiso de micrófono limitado","Cámara autorizada solo si corresponde","Pantalla autorizada solo si corresponde","Notificaciones y ventanas sensibles cerradas","Consentimiento de terceros obtenido","Datos confidenciales retirados","Idioma y contexto definidos","Interrupción de emergencia conocida","Transcripción revisada","Fuentes y datos importantes validados","Micrófono, cámara y pantalla detenidos","Retención y eliminación revisadas","Resultado y observaciones registrados"]
grid_pdf(OUT/"lista-verificacion-audio-privacidad-permisos.pdf","Lista de verificación de audio, privacidad y permisos",["N°","Control","Cumple","No cumple","N/A","Observaciones"],[[i+1,x,"","","",""] for i,x in enumerate(checks)],[24,245,42,48,34,125])
rub=[("Definición del objetivo",15,"Objetivo, audiencia y éxito inequívocos"),("Preparación técnica y permisos",15,"Dispositivo, entorno y permisos verificados"),("Estructura y continuidad",15,"Diálogo ordenado y contextual"),("Interrupciones y correcciones",15,"Redirige, aclara y corrige"),("Aplicación profesional o educativa",10,"Uso pertinente y transferible"),("Privacidad y consentimiento",10,"Protección y autorización demostrables"),("Validación del resultado",10,"Datos, resumen y transcripción revisados"),("Registro y reflexión",10,"Trazabilidad y mejoras documentadas")]
grid_pdf(OUT/"rubrica-laboratorio-live.pdf","Rúbrica del Laboratorio ChatGPT Voice — experiencia Live",["Criterio","Máx.","Excelente","Competente","En desarrollo","Requiere revisión","Obtenido"],[[a,b,c,"Cumple lo esencial","Cumplimiento parcial","Debe repetirse",""] for a,b,c in rub],[82,28,105,82,82,82,42])
headers=["ID","Fecha","Objetivo","Opción","Dispositivo","Permisos","Consentimiento","Resultado","Validación","Incidencias","Próximo paso","Responsable"]
grid_pdf(OUT/"registro-sesiones-live.pdf","Registro de sesiones, objetivos, resultados y observaciones",headers,[[i+1,"","","","","","","","","","",""] for i in range(20)],[22,38,60,42,48,48,48,58,54,54,55,50])
