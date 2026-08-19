from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

OUT=Path(r"C:\Users\ASUS\Codex_Proyectos\CIVG 2026\phase1c_output")
OUT.mkdir(parents=True,exist_ok=True)
NAVY="07152B"; BLUE="0B3F75"; GOLD="F0A22E"; LIGHT="EEF3F8"

def shade(cell,fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement("w:shd"); shd.set(qn("w:fill"),fill); tcPr.append(shd)

def base_doc(title,subtitle):
    d=Document(); s=d.sections[0]; s.top_margin=s.bottom_margin=s.left_margin=s.right_margin=Inches(1)
    styles=d.styles
    normal=styles["Normal"]; normal.font.name="Calibri"; normal.font.size=Pt(11); normal.font.color.rgb=RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.25
    for name,size,color,b,a in [("Heading 1",16,BLUE,18,10),("Heading 2",13,BLUE,14,7),("Heading 3",12,NAVY,10,5)]:
        st=styles[name]; st.font.name="Calibri"; st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(color); st.paragraph_format.space_before=Pt(b); st.paragraph_format.space_after=Pt(a)
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(title); r.bold=True; r.font.size=Pt(24); r.font.color.rgb=RGBColor.from_string(GOLD)
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(subtitle); r.font.size=Pt(10); r.font.color.rgb=RGBColor.from_string(BLUE)
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("ZYRON AI LAB · Versión 2026-07-26"); r.bold=True; r.font.size=Pt(9)
    return d

def add_form(d,sections):
    for heading,fields in sections:
        d.add_heading(heading,level=1)
        t=d.add_table(rows=1,cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False; t.columns[0].width=Inches(1.9); t.columns[1].width=Inches(4.6)
        t.rows[0].cells[0].text="Campo"; t.rows[0].cells[1].text="Contenido"
        for c in t.rows[0].cells:
            shade(c,BLUE); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for run in c.paragraphs[0].runs:
                run.bold=True; run.font.color.rgb=RGBColor(255,255,255)
        for f in fields:
            cells=t.add_row().cells; cells[0].text=f; cells[1].text="\n"; shade(cells[0],LIGHT)

planning=[
("1. Definición del encargo",["Nombre de la tarea","Necesidad o problema","Objetivo","Audiencia","Alcance incluido","Exclusiones","Fecha requerida","Responsable de aprobación"]),
("2. Contexto y fuentes",["Archivos disponibles","Fuentes oficiales","Aplicaciones conectadas","Permisos autorizados","Datos sensibles excluidos","Preguntas de aclaración"]),
("3. Plan de trabajo",["Etapas","Dependencias","Responsables","Puntos de aprobación","Entregable por etapa","Criterios de aceptación"]),
("4. Entrega y control",["Formatos finales","Ubicación de archivos","Validación de fórmulas y datos","Validación visual","Registro de versiones","Aprobación final","Observaciones"])
]
d=base_doc("Plantilla de planificación de tareas complejas","Necesidad → contexto → plan → aprobación → ejecución → validación → entrega")
add_form(d,planning); d.save(OUT/"plantilla-planificacion-work.docx")

security=[
("1. Clasificación de la información","Identificar información pública, interna, confidencial, personal o restringida antes de compartirla."),
("2. Acceso mínimo","Habilitar solo las aplicaciones, archivos y acciones necesarios para el encargo."),
("3. Conectores autorizados","Confirmar cuenta, organización, permisos, finalidad y responsable antes de conectar una aplicación."),
("4. Datos personales","Minimizar datos, justificar su uso y aplicar las políticas internas y legales correspondientes."),
("5. Credenciales y secretos","No incluir contraseñas, tokens, claves API, códigos de recuperación ni secretos en archivos o instrucciones."),
("6. Archivos y metadatos","Revisar comentarios, propiedades, historial, autores y contenido oculto antes de compartir."),
("7. Fuentes y derechos","Comprobar vigencia, autoría, licencias, atribución y autorización de publicación."),
("8. Fórmulas y cálculos","Validar entradas, referencias, supuestos, unidades, fórmulas y resultados."),
("9. Formatos y diseño","Abrir cada archivo final y revisar paginación, tablas, gráficos, tipografía y elementos recortados."),
("10. Versiones y almacenamiento","Registrar nombre, versión, ubicación, responsable, copia de respaldo y retención."),
("11. Aprobación humana","No enviar, publicar ni utilizar para una decisión importante sin una revisión competente."),
("12. Respuesta ante errores","Detener la distribución, conservar evidencia, corregir, volver a validar y comunicar la versión afectada."),
("13. Disponibilidad y superficies","Confirmar plan, plataforma, aplicaciones y formatos antes de prometer una capacidad."),
("14. Lista final","Revisar audiencia, acceso, datos, fuentes, fórmulas, diseño, versiones, destinatarios y aprobación.")
]
d=base_doc("Guía de seguridad y manejo de información","Controles prácticos para trabajos profesionales con ChatGPT Work")
for h,b in security: d.add_heading(h,level=1); d.add_paragraph(b)
d.add_paragraph("Nota: esta guía es educativa y no sustituye políticas internas, asesoría jurídica, controles de seguridad ni revisión profesional.")
d.save(OUT/"guia-seguridad-informacion-work.docx")

styles=getSampleStyleSheet()
title_style=ParagraphStyle("T",parent=styles["Title"],fontName="Helvetica-Bold",fontSize=20,textColor=colors.HexColor("#"+GOLD),alignment=TA_CENTER,spaceAfter=10)
sub_style=ParagraphStyle("S",parent=styles["BodyText"],fontName="Helvetica",fontSize=9,textColor=colors.HexColor("#"+BLUE),alignment=TA_CENTER,spaceAfter=10)
body=ParagraphStyle("B",parent=styles["BodyText"],fontName="Helvetica",fontSize=7.8,leading=9.3,textColor=colors.HexColor("#"+NAVY))
head=ParagraphStyle("H",parent=body,fontName="Helvetica-Bold",textColor=colors.white,alignment=TA_CENTER)

def pdf_header(title): return [Paragraph(title,title_style),Paragraph("Laboratorio Work · ZYRON AI LAB · Versión 2026-07-26",sub_style)]
def grid_pdf(name,title,headers,rows,widths,pagesize=landscape(letter)):
    story=pdf_header(title)
    data=[[Paragraph(f"<b>{x}</b>",head) for x in headers]]+[[Paragraph(str(x),body) for x in row] for row in rows]
    t=Table(data,colWidths=widths,repeatRows=1)
    t.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor("#"+BLUE)),("TEXTCOLOR",(0,0),(-1,0),colors.white),("GRID",(0,0),(-1,-1),.35,colors.HexColor("#AAB8C5")),("VALIGN",(0,0),(-1,-1),"TOP"),("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white,colors.HexColor("#"+LIGHT)]),("LEFTPADDING",(0,0),(-1,-1),4),("RIGHTPADDING",(0,0),(-1,-1),4),("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4)]))
    story.append(t); SimpleDocTemplate(str(OUT/name),pagesize=pagesize,rightMargin=24,leftMargin=24,topMargin=26,bottomMargin=26).build(story)

plan_fields=[f for _,fs in planning for f in fs]
grid_pdf("plantilla-planificacion-work.pdf","Plantilla de planificación de tareas complejas",["Campo","Contenido"],[[x,""] for x in plan_fields],[160,360],letter)
story=pdf_header("Guía de seguridad y manejo de información")
for h,b in security: story += [Paragraph(f"<b>{h}</b>",ParagraphStyle("sec"+str(len(story)),parent=body,fontSize=9.2,leading=10.5,textColor=colors.HexColor("#"+BLUE),spaceBefore=3,spaceAfter=1)),Paragraph(b,body)]
story.append(Spacer(1,4)); story.append(Paragraph("<b>Nota:</b> esta guía es educativa y no sustituye políticas internas, asesoría jurídica, controles de seguridad ni revisión profesional.",body))
SimpleDocTemplate(str(OUT/"guia-seguridad-informacion-work.pdf"),pagesize=letter,rightMargin=32,leftMargin=32,topMargin=28,bottomMargin=28).build(story)

checks=["Objetivo y audiencia definidos","Alcance y exclusiones confirmados","Preguntas de aclaración resueltas","Plan aprobado por una persona","Fuentes identificadas y vigentes","Archivos autorizados","Permisos mínimos aplicados","Datos sensibles excluidos o protegidos","Etapas y responsables registrados","Entregables intermedios revisados","Fórmulas y cálculos comprobados","Datos y unidades verificados","Enlaces y referencias correctos","Diseño y formato inspeccionados","Metadatos y comentarios revisados","Versiones y ubicación registradas","Formatos finales abiertos","Destinatarios confirmados","Copia de respaldo disponible","Aprobación final registrada"]
grid_pdf("lista-verificacion-work.pdf","Lista de verificación para trabajos con Work",["N°","Criterio","Cumple","No cumple","N/A","Observaciones"],[[i+1,x,"","","",""] for i,x in enumerate(checks)],[28,220,55,60,45,300])
rub=[("Definición del encargo",15),("Planificación de etapas",15),("Investigación y fuentes",15),("Calidad de entregables",15),("Datos y fórmulas",10),("Trazabilidad",10),("Seguridad y permisos",10),("Validación profesional",10)]
grid_pdf("rubrica-laboratorio-work.pdf","Rúbrica de evaluación del Laboratorio Work",["Criterio","Máx.","Excelente","Competente","En desarrollo","Requiere revisión","Obtenido"],[[a,b,"Dominio completo","Cumple lo esencial","Cumplimiento parcial","Debe rehacerse",""] for a,b in rub],[130,40,120,120,120,120,55])
headers=["ID","Etapa","Tarea","Fuente o archivo","Herramienta","Responsable","Estado","Entregable","Validación","Incidencia","Versión","Aprobación"]
grid_pdf("registro-tareas-entregables-work.pdf","Registro de tareas, etapas, fuentes y entregables",headers,[["","","","","","","","","","","",""] for _ in range(10)],[32,58,70,78,62,62,52,70,65,65,45,55])
