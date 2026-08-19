from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
import sys

OUT = Path(sys.argv[1]); OUT.mkdir(parents=True, exist_ok=True)
NAVY="#07182F"; BLUE="#0B3B70"; GOLD="#EBA03D"; CYAN="#00D4FF"; LIGHT="#F4F7FA"

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement("w:shd"); shd.set(qn("w:fill"), fill.replace("#","")); tcPr.append(shd)

def base_doc(title, subtitle):
    d=Document(); sec=d.sections[0]; sec.top_margin=Inches(.7); sec.bottom_margin=Inches(.7); sec.left_margin=Inches(.8); sec.right_margin=Inches(.8)
    styles=d.styles
    normal=styles["Normal"]; normal.font.name="Calibri"; normal.font.size=Pt(10.5); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.15
    for n,size,color in [("Heading 1",16,GOLD),("Heading 2",13,BLUE),("Heading 3",11.5,BLUE)]:
        s=styles[n]; s.font.name="Calibri"; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color.replace("#","")); s.paragraph_format.space_before=Pt(10); s.paragraph_format.space_after=Pt(5)
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(title); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=RGBColor.from_string(GOLD.replace("#",""))
    p=d.add_paragraph(subtitle); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].italic=True; p.runs[0].font.color.rgb=RGBColor.from_string(BLUE.replace("#",""))
    d.add_paragraph("ZYRON AI LAB · Versión 2026-07-26").alignment=WD_ALIGN_PARAGRAPH.CENTER
    return d

def add_field_table(d, fields):
    t=d.add_table(rows=1, cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    t.columns[0].width=Inches(1.8); t.columns[1].width=Inches(4.7)
    h=t.rows[0].cells; h[0].merge(h[1]); h[0].text="REGISTRO EDITABLE"; shade(h[0], BLUE)
    for run in h[0].paragraphs[0].runs: run.font.color.rgb=RGBColor(255,255,255); run.bold=True
    for label, hint in fields:
        c=t.add_row().cells; c[0].text=label; c[1].text=hint
        shade(c[0], "E8EEF5"); c[0].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        c[0].paragraphs[0].runs[0].bold=True
        c[1].paragraphs[0].paragraph_format.space_after=Pt(9 if len(hint)<30 else 13)
    return t

prompt_fields=[
("Nombre del proyecto",""),("Organización o responsable",""),("Fecha","AAAA-MM-DD"),("Objetivo del sitio",""),
("Público objetivo",""),("Problema que resolverá",""),("Mensaje principal",""),("Secciones requeridas",""),
("Contenido disponible",""),("Identidad visual",""),("Colores",""),("Tipografías",""),
("Logotipo e imágenes",""),("Funciones requeridas",""),("Formularios",""),("Navegación",""),
("Dispositivos prioritarios",""),("Accesibilidad",""),("Dominio previsto",""),("Publicación",""),
("Mantenimiento",""),("Responsable de aprobación",""),("Observaciones","")
]
d=base_doc("Brief estructurado para crear un Site","Formulario editable para definir propósito, audiencia, contenidos, funciones y criterios de publicación.")
add_field_table(d,prompt_fields); d.save(OUT/"brief-estructurado-sites.docx")

security_sections=[
("1. Revisión del contenido","Comprobar exactitud, vigencia, fuentes, tono y autorización de todos los textos, imágenes, archivos y datos."),
("2. Protección de datos personales","Aplicar minimización, autorización, finalidad definida y controles acordes con la normativa y las políticas internas."),
("3. Formularios y consentimiento","Recoger solo los datos necesarios, explicar su finalidad y obtener el consentimiento aplicable antes del envío."),
("4. Contraseñas y credenciales","Nunca compartir contraseñas, tokens, claves API, códigos de recuperación ni secretos de acceso."),
("5. Dominios y alojamiento","Verificar propiedad del dominio, registros DNS, audiencia, permisos y configuración del alojamiento antes de publicar."),
("6. Copias de respaldo","Guardar una versión revisada y una copia recuperable de los archivos antes de cada despliegue."),
("7. Accesibilidad","Revisar navegación por teclado, estructura, contraste, texto alternativo, etiquetas y visualización responsive."),
("8. Derechos de imágenes y textos","Confirmar autoría, licencias, atribuciones y autorización de publicación."),
("9. Enlaces externos","Comprobar destino, vigencia, seguridad y apertura adecuada de cada enlace externo."),
("10. Actualizaciones y mantenimiento","Definir responsable, frecuencia de revisión, registro de cambios y procedimiento de nueva publicación."),
("11. Protección contra spam","Aplicar controles razonables a formularios y canales de contacto sin recopilar datos innecesarios."),
("12. Procedimiento ante errores","Retirar o limitar la versión afectada, conservar evidencia, corregir, validar y volver a publicar solo con aprobación."),
("13. Responsable de publicación","Identificar a la persona que autoriza audiencia, versión y fecha de publicación."),
("14. Lista final antes de publicar","Confirmar contenido, acceso, privacidad, formularios, enlaces, archivos, dispositivos, respaldo y aprobación.")
]
d=base_doc("Guía de publicación segura de Sites","Referencia educativa para trabajar con Sites de manera supervisada.")
d.sections[0].top_margin=Inches(.45); d.sections[0].bottom_margin=Inches(.45)
d.styles["Normal"].font.size=Pt(9); d.styles["Normal"].paragraph_format.space_after=Pt(2); d.styles["Normal"].paragraph_format.line_spacing=1
d.styles["Heading 1"].font.size=Pt(11); d.styles["Heading 1"].paragraph_format.space_before=Pt(5); d.styles["Heading 1"].paragraph_format.space_after=Pt(2)
for h,b in security_sections: d.add_heading(h,level=1); d.add_paragraph(b)
p=d.add_paragraph(); r=p.add_run("Nota: "); r.bold=True; p.add_run("Esta guía es educativa y no sustituye las políticas internas, la asesoría jurídica ni los protocolos de seguridad de cada organización.")
d.save(OUT/"guia-publicacion-segura-sites.docx")

styles=getSampleStyleSheet()
title_style=ParagraphStyle("ZTitle",parent=styles["Title"],fontName="Helvetica-Bold",fontSize=20,textColor=colors.HexColor(GOLD),alignment=TA_CENTER,spaceAfter=12)
h_style=ParagraphStyle("ZH",parent=styles["Heading2"],fontName="Helvetica-Bold",fontSize=12,textColor=colors.HexColor(BLUE),spaceBefore=8,spaceAfter=4)
body_style=ParagraphStyle("ZBody",parent=styles["BodyText"],fontName="Helvetica",fontSize=9.3,leading=12,textColor=colors.HexColor(NAVY))
small_style=ParagraphStyle("ZSmall",parent=body_style,fontSize=7,leading=8)
header_style=ParagraphStyle("ZHeader",parent=small_style,textColor=colors.white,alignment=TA_CENTER)

def pdf_header(title, subtitle):
    return [Paragraph(title,title_style),Paragraph(subtitle,ParagraphStyle("sub",parent=body_style,alignment=TA_CENTER,textColor=colors.HexColor(BLUE))),Spacer(1,10)]

def make_prompt_pdf():
    story=pdf_header("Brief estructurado para crear un Site","Laboratorio Sites · ZYRON AI LAB · Versión 2026-07-26")
    data=[[Paragraph("<b>Campo</b>",header_style),Paragraph("<b>Contenido</b>",header_style)]]
    for a,b in prompt_fields:data.append([Paragraph(a,small_style),Paragraph((b+"<br/><br/>") if b else "<br/><br/>",small_style)])
    t=Table(data,colWidths=[1.65*72,4.85*72],repeatRows=1)
    t.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor(BLUE)),("TEXTCOLOR",(0,0),(-1,0),colors.white),("BACKGROUND",(0,1),(0,-1),colors.HexColor("#E8EEF5")),("GRID",(0,0),(-1,-1),.4,colors.HexColor("#AAB8C5")),("VALIGN",(0,0),(-1,-1),"TOP"),("LEFTPADDING",(0,0),(-1,-1),4),("RIGHTPADDING",(0,0),(-1,-1),4),("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3)]))
    story.append(t); SimpleDocTemplate(str(OUT/"brief-estructurado-sites.pdf"),pagesize=letter,rightMargin=24,leftMargin=24,topMargin=24,bottomMargin=24).build(story)

def make_security_pdf():
    story=pdf_header("Guía de publicación segura de Sites","Laboratorio Sites · ZYRON AI LAB · Versión 2026-07-26")
    compact_h=ParagraphStyle("CompactH",parent=h_style,fontSize=10,spaceBefore=4,spaceAfter=2)
    compact_b=ParagraphStyle("CompactB",parent=body_style,fontSize=8.2,leading=9.4)
    for h,b in security_sections: story += [Paragraph(h,compact_h),Paragraph(b,compact_b)]
    story += [Spacer(1,5),Paragraph("<b>Nota:</b> Esta guía es educativa y no sustituye las políticas internas, la asesoría jurídica ni los protocolos de seguridad de cada organización.",compact_b)]
    SimpleDocTemplate(str(OUT/"guia-publicacion-segura-sites.pdf"),pagesize=letter,rightMargin=32,leftMargin=32,topMargin=28,bottomMargin=28).build(story)

def make_grid_pdf(filename,title,headers,rows,widths,orientation=landscape(letter)):
    story=pdf_header(title,"Laboratorio Sites · ZYRON AI LAB · Versión 2026-07-26")
    data=[[Paragraph(f"<b>{x}</b>",header_style) for x in headers]]+[[Paragraph(str(x),small_style) for x in row] for row in rows]
    t=Table(data,colWidths=widths,repeatRows=1)
    t.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor(BLUE)),("TEXTCOLOR",(0,0),(-1,0),colors.white),("GRID",(0,0),(-1,-1),.35,colors.HexColor("#AAB8C5")),("VALIGN",(0,0),(-1,-1),"TOP"),("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white,colors.HexColor(LIGHT)]),("LEFTPADDING",(0,0),(-1,-1),4),("RIGHTPADDING",(0,0),(-1,-1),4),("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5)]))
    story.append(t); SimpleDocTemplate(str(OUT/filename),pagesize=orientation,rightMargin=24,leftMargin=24,topMargin=30,bottomMargin=30).build(story)

criteria=["Objetivo del sitio definido","Público identificado","Contenido aprobado","Navegación funcional","Enlaces internos correctos","Enlaces externos seguros","Imágenes disponibles","Texto alternativo","Contraste legible","Visualización móvil","Formularios probados","Política de privacidad","Consentimiento de datos","Dominio y alojamiento","Copia de respaldo","Responsable de aprobación","Fecha de publicación","Plan de mantenimiento"]
make_grid_pdf("lista-verificacion-publicacion-sites.pdf","Lista de verificación para publicar un Site",["N°","Criterio","Cumple","No cumple","No aplica","Observaciones"],[[i+1,c,"","","",""] for i,c in enumerate(criteria)],[28,210,55,62,55,310])

rubrics=[("Definición del propósito",15),("Arquitectura y navegación",15),("Calidad del contenido",15),("Diseño visual y coherencia",15),("Adaptación móvil",10),("Accesibilidad",10),("Funcionalidad y enlaces",10),("Seguridad, privacidad y publicación",10)]
make_grid_pdf("rubrica-laboratorio-sites.pdf","Rúbrica de evaluación del laboratorio",["Criterio","Máx.","Excelente","Competente","En desarrollo","Requiere revisión","Obtenido"],[[a,b,"Dominio completo","Cumple lo esencial","Cumplimiento parcial","Debe rehacerse",""] for a,b in rubrics],[120,36,125,125,125,125,50])

iteration_headers=["N° de versión","Fecha","Página o sección","Cambio realizado","Motivo","Archivo modificado","Responsable","Revisión realizada","Estado","Fecha de aprobación","Observaciones"]
make_grid_pdf("registro-versiones-sites.pdf","Registro de versiones y cambios del Site",iteration_headers,[["","","","","","","","","","",""] for _ in range(8)],[48,45,70,82,66,75,60,72,48,62,72])
make_prompt_pdf(); make_security_pdf()
