from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

ROOT=Path(r"C:\Users\ASUS\Desktop\Empresa Zyron\ZYRON PRIMUS")
OUT=ROOT/"recursos"/"laboratorio-computer-use"; OUT.mkdir(parents=True,exist_ok=True)
NAVY="07152B"; BLUE="0B3F75"; GOLD="F0A22E"; LIGHT="EEF3F8"

def shade(cell,fill):
 p=cell._tc.get_or_add_tcPr(); x=OxmlElement("w:shd"); x.set(qn("w:fill"),fill); p.append(x)
def prevent_split(row,header=False):
 p=row._tr.get_or_add_trPr(); p.append(OxmlElement("w:cantSplit"))
 if header:
  x=OxmlElement("w:tblHeader"); x.set(qn("w:val"),"true"); p.append(x)
def base(title,subtitle):
 d=Document(); s=d.sections[0]; s.top_margin=s.bottom_margin=s.left_margin=s.right_margin=Inches(1)
 n=d.styles["Normal"]; n.font.name="Calibri"; n.font.size=Pt(11); n.font.color.rgb=RGBColor.from_string(NAVY); n.paragraph_format.space_after=Pt(6); n.paragraph_format.line_spacing=1.25
 for k,z,c,b,a in [("Heading 1",16,BLUE,18,10),("Heading 2",13,BLUE,14,7),("Heading 3",12,NAVY,10,5)]:
  st=d.styles[k]; st.font.name="Calibri"; st.font.size=Pt(z); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(c); st.paragraph_format.space_before=Pt(b); st.paragraph_format.space_after=Pt(a)
 p=d.add_paragraph();p.alignment=1;r=p.add_run(title);r.bold=True;r.font.size=Pt(24);r.font.color.rgb=RGBColor.from_string(GOLD)
 p=d.add_paragraph();p.alignment=1;p.add_run(subtitle).italic=True
 p=d.add_paragraph();p.alignment=1;p.add_run("ZYRON AI LAB · Revisión 26 de julio de 2026").bold=True
 return d
def grid(d,headers,rows,widths):
 t=d.add_table(rows=1,cols=len(headers));t.alignment=WD_TABLE_ALIGNMENT.CENTER;t.autofit=False
 for i,h in enumerate(headers):
  t.rows[0].cells[i].text=h;shade(t.rows[0].cells[i],BLUE)
  for r in t.rows[0].cells[i].paragraphs[0].runs:r.bold=True;r.font.color.rgb=RGBColor(255,255,255)
 prevent_split(t.rows[0],True)
 for row in rows:
  cells=t.add_row().cells
  for i,v in enumerate(row):cells[i].text=str(v);shade(cells[i],LIGHT if len(t.rows)%2==0 else "FFFFFF")
  prevent_split(t.rows[-1])
 for row in t.rows:
  for i,w in enumerate(widths):row.cells[i].width=Inches(w)
 return t

fields=[
("1. Resultado y alcance",["Objetivo verificable","Resultado esperado","Aplicaciones o sitios autorizados","Exclusiones","Criterio de éxito","Responsable de aprobación"]),
("2. Datos y permisos",["Datos permitidos","Datos prohibidos","Cuenta autorizada","Permisos mínimos","Credenciales: toma de control obligatoria","Retención y eliminación"]),
("3. Acciones",["Pasos previstos","Acciones reversibles","Acciones sensibles","Acciones prohibidas","Puntos de confirmación","Procedimiento de cancelación"]),
("4. Validación",["Evidencia mínima","Resultado obtenido","Incidencias","Correcciones","Aprobación humana","Cierre de sesión y permisos"])
]
d=base("Plantilla de planificación de tareas con Computer Use","Objetivo, permisos, confirmaciones, ejecución y validación")
for h,items in fields:
 d.add_heading(h,1);grid(d,["Campo","Contenido"],[[x,""] for x in items],[2.1,4.4])
d.add_heading("Regla de uso",1);d.add_paragraph("No utilice esta plantilla para autorizar pagos, compras, publicaciones, eliminaciones ni decisiones de alto impacto sin revisión y confirmación humana competente.")
d.save(OUT/"plantilla-planificacion-computer-use.docx")

security=[
("1. Delimitar la tarea","Definir resultado, sitios, aplicaciones, cuentas, datos permitidos y acciones expresamente excluidas."),
("2. Aplicar acceso mínimo","Autorizar solo las superficies y archivos necesarios; cerrar sesiones y retirar permisos al terminar."),
("3. Proteger credenciales","No escribir contraseñas, códigos, tokens ni secretos en el chat. Tomar el control en el navegador o aplicación autorizada."),
("4. Clasificar acciones","Separar acciones reversibles de envíos, publicaciones, compras, pagos, cambios de cuenta, eliminación y otros efectos externos."),
("5. Confirmar antes del efecto","Revisar destinatario, cuenta, importe, archivo, texto y consecuencia antes de cualquier acción sensible."),
("6. Tratar la interfaz como no confiable","Las páginas pueden contener prompt injection, phishing, anuncios engañosos o instrucciones que intenten desviar al agente."),
("7. Supervisar","Detener la tarea si aparece otra cuenta, un sitio inesperado, información sensible, un CAPTCHA o una acción fuera del alcance."),
("8. Preferir integración estructurada","Usar API o conector cuando exista una vía autorizada, estable y auditable para el mismo objetivo."),
("9. Validar el resultado","Comparar el estado final con el criterio de éxito y revisar directamente los cambios producidos."),
("10. Registrar","Documentar pasos, confirmaciones, incidencias, correcciones, responsable y aprobación final."),
("11. Responder a incidentes","Detener, retirar acceso, conservar evidencia mínima, informar, corregir y revisar la exposición de datos."),
("12. Responsabilidad humana","Computer Use no sustituye decisiones profesionales, controles organizacionales ni responsabilidad humana.")
]
d=base("Guía de seguridad y supervisión humana","Controles para agentes que interactúan con interfaces digitales")
for h,b in security:d.add_heading(h,1);d.add_paragraph(b)
d.add_heading("Lista de cierre",1)
for x in ["Tarea detenida o finalizada","Cuenta correcta verificada","Acciones sensibles confirmadas","Resultado revisado","Incidencias registradas","Permisos retirados","Datos y sesiones tratados"]:d.add_paragraph(x,style="List Bullet")
d.add_paragraph("Material educativo: no sustituye políticas internas, asesoría jurídica, controles de seguridad ni revisión profesional.")
d.save(OUT/"guia-seguridad-supervision-computer-use.docx")

styles=getSampleStyleSheet(); T=ParagraphStyle("T",parent=styles["Title"],fontName="Helvetica-Bold",fontSize=19,textColor=colors.HexColor("#"+GOLD),alignment=1,spaceAfter=8); B=ParagraphStyle("B",parent=styles["BodyText"],fontName="Helvetica",fontSize=7.5,leading=9,textColor=colors.HexColor("#"+NAVY)); H=ParagraphStyle("H",parent=B,fontName="Helvetica-Bold",textColor=colors.white,alignment=1)
def grid_pdf(name,title,headers,rows,widths,pagesize=letter):
 st=[Paragraph(title,T),Paragraph("ZYRON AI LAB · 26-07-2026",ParagraphStyle("S",parent=B,alignment=1)),Spacer(1,6)]
 data=[[Paragraph(str(x),H) for x in headers]]+[[Paragraph(str(x),B) for x in row] for row in rows]
 t=Table(data,colWidths=widths,repeatRows=1,splitByRow=1);t.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor("#"+BLUE)),("GRID",(0,0),(-1,-1),.35,colors.HexColor("#AAB8C5")),("VALIGN",(0,0),(-1,-1),"TOP"),("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white,colors.HexColor("#"+LIGHT)]),("LEFTPADDING",(0,0),(-1,-1),3),("RIGHTPADDING",(0,0),(-1,-1),3),("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3)]));st.append(t)
 SimpleDocTemplate(str(OUT/name),pagesize=pagesize,rightMargin=24,leftMargin=24,topMargin=28,bottomMargin=28).build(st)

allfields=[x for _,items in fields for x in items]
grid_pdf("plantilla-planificacion-computer-use.pdf","Plantilla de planificación de tareas con Computer Use",["Campo","Contenido"],[[x,""] for x in allfields],[160,360])
grid_pdf("guia-seguridad-supervision-computer-use.pdf","Guía de seguridad y supervisión humana",["Control","Aplicación"],security,[145,375])
checks=["Objetivo y criterio de éxito definidos","Sitios y aplicaciones autorizados","Cuenta correcta verificada","Datos permitidos identificados","Datos sensibles excluidos","Permisos mínimos aplicados","Credenciales reservadas a toma de control","Acciones sensibles clasificadas","Acciones prohibidas registradas","Confirmación previa definida","Procedimiento de cancelación conocido","Prompt injection considerada","Sitios y destinatarios verificados","Compras y pagos excluidos","Envíos y publicaciones detenidos para revisión","Eliminaciones detenidas para revisión","Capturas limitadas a evidencia necesaria","Resultado comparado con objetivo","Incidencias y correcciones registradas","Permisos y sesiones cerrados"]
grid_pdf("lista-verificacion-permisos-acciones-sensibles.pdf","Lista de verificación de permisos y acciones sensibles",["N°","Control","Cumple","No cumple","N/A","Observaciones"],[[i+1,x,"","","",""] for i,x in enumerate(checks)],[26,245,42,50,35,125],landscape(letter))
rub=[("Objetivo y alcance",15),("Entorno y permisos",15),("Acciones sensibles",15),("Supervisión y confirmaciones",15),("Ejecución y registro",10),("Seguridad y privacidad",10),("Validación",10),("Informe y evidencias",10)]
grid_pdf("rubrica-laboratorio-computer-use.pdf","Rúbrica del Laboratorio Computer Use",["Criterio","Máx.","Excelente","Competente","En desarrollo","Requiere revisión","Obtenido"],[[a,b,"Dominio completo","Cumple lo esencial","Cumplimiento parcial","Debe repetirse",""] for a,b in rub],[110,35,110,110,110,110,55],landscape(letter))
headers=["ID","Fecha","Paso","Interfaz","Acción","Confirmación","Responsable","Resultado","Incidencia","Corrección","Evidencia","Aprobación"]
grid_pdf("registro-acciones-computer-use.pdf","Registro de acciones, confirmaciones, incidencias y resultados",headers,[[i+1,"","","","","","","","","","",""] for i in range(15)],[25,45,55,55,70,60,60,65,65,65,55,55],landscape(letter))
