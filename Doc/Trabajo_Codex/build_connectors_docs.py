from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

ROOT=Path(r"C:\Users\ASUS\Desktop\Empresa Zyron\ZYRON PRIMUS")
OUT=ROOT/"recursos"/"laboratorio-conectores"; OUT.mkdir(parents=True,exist_ok=True)
NAVY="07152B"; BLUE="0B3F75"; GOLD="F0A22E"; LIGHT="EEF3F8"
def shade(cell,fill):
 p=cell._tc.get_or_add_tcPr(); x=OxmlElement("w:shd"); x.set(qn("w:fill"),fill); p.append(x)
def row_rule(row,header=False):
 p=row._tr.get_or_add_trPr(); p.append(OxmlElement("w:cantSplit"))
 if header:
  x=OxmlElement("w:tblHeader"); x.set(qn("w:val"),"true"); p.append(x)
def base(title,subtitle):
 d=Document(); s=d.sections[0]; s.top_margin=s.bottom_margin=s.left_margin=s.right_margin=Inches(1)
 n=d.styles["Normal"]; n.font.name="Calibri"; n.font.size=Pt(11); n.font.color.rgb=RGBColor.from_string(NAVY); n.paragraph_format.space_after=Pt(6); n.paragraph_format.line_spacing=1.25
 for k,z,c,b,a in [("Heading 1",16,BLUE,18,10),("Heading 2",13,BLUE,14,7),("Heading 3",12,NAVY,10,5)]:
  st=d.styles[k]; st.font.name="Calibri"; st.font.size=Pt(z); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(c); st.paragraph_format.space_before=Pt(b); st.paragraph_format.space_after=Pt(a)
 p=d.add_paragraph(); p.alignment=1; r=p.add_run(title); r.bold=True; r.font.size=Pt(24); r.font.color.rgb=RGBColor.from_string(GOLD)
 p=d.add_paragraph(); p.alignment=1; p.add_run(subtitle).italic=True
 p=d.add_paragraph(); p.alignment=1; p.add_run("ZYRON AI LAB · Revisión 28 de julio de 2026").bold=True
 return d
def grid(d,headers,rows,widths):
 t=d.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
 for i,h in enumerate(headers):
  t.rows[0].cells[i].text=h; shade(t.rows[0].cells[i],BLUE)
  for r in t.rows[0].cells[i].paragraphs[0].runs:r.bold=True;r.font.color.rgb=RGBColor(255,255,255)
 row_rule(t.rows[0],True)
 for row in rows:
  cells=t.add_row().cells
  for i,v in enumerate(row): cells[i].text=str(v); shade(cells[i],LIGHT if len(t.rows)%2==0 else "FFFFFF")
  row_rule(t.rows[-1])
 for row in t.rows:
  for i,w in enumerate(widths): row.cells[i].width=Inches(w)
 return t

sections=[
("1. Objetivo y alcance",["Objetivo verificable","Resultado esperado","Fuente o servicio","Cuenta autorizada","Datos necesarios","Datos excluidos"]),
("2. Autorización",["Disponibilidad confirmada","Permisos solicitados","Permisos mínimos aprobados","Responsable de autorización","Fecha de revisión","Procedimiento de revocación"]),
("3. Consultas y acciones",["Consulta prevista","Resultados esperados","Acciones de lectura","Acciones de escritura","Confirmaciones requeridas","Acciones prohibidas"]),
("4. Validación y cierre",["Fuente verificada","Resultado validado","Evidencia mínima","Incidencias","Decisión de mantener o revocar","Aprobación humana"])
]
d=base("Plantilla de planificación y autorización de conectores","Objetivo, fuente, permisos, consultas, acciones y cierre")
for h,items in sections:
 d.add_heading(h,1); grid(d,["Campo","Contenido"],[[x,""] for x in items],[2.1,4.4])
d.add_heading("Regla de uso",1); d.add_paragraph("Conectar una fuente no autoriza acceso ilimitado ni acciones externas. Conceda permisos mínimos, verifique la cuenta y confirme por separado cualquier creación, edición, envío o eliminación.")
d.save(OUT/"plantilla-planificacion-autorizacion-conectores.docx")

security=[
("1. Confirmar la cuenta","Verificar identidad, espacio de trabajo y servicio antes de autorizar."),
("2. Aplicar permisos mínimos","Conceder solo lectura, búsqueda o acciones estrictamente necesarias."),
("3. Limitar los datos","Excluir secretos, credenciales, categorías sensibles y contenido ajeno al objetivo."),
("4. Separar consulta y acción","Distinguir búsqueda, lectura, creación, edición, envío y eliminación."),
("5. Confirmar efectos externos","Revisar destinatario, fecha, archivo, contenido y consecuencia antes de escribir."),
("6. Verificar la fuente","Abrir el origen y comprobar actualidad, propiedad y contexto del resultado."),
("7. Controlar retención y sync","Revisar las opciones reales del servicio; conectar no significa copiar toda la cuenta."),
("8. Registrar el uso","Documentar fuente, permisos, consulta, resultados, acciones, incidencias y aprobación."),
("9. Revocar cuando termine","Desconectar o retirar permisos si el acceso ya no es necesario."),
("10. Responder a incidentes","Detener, revocar, conservar evidencia mínima, informar y revisar exposición."),
("11. Respetar gobierno institucional","Aplicar políticas internas, roles administrativos y requisitos regulatorios."),
("12. Mantener supervisión humana","Los conectores no sustituyen responsabilidad profesional ni revisión competente.")
]
d=base("Guía de seguridad, privacidad y revocación","Controles para fuentes y servicios conectados")
for h,b in security:d.add_heading(h,1);d.add_paragraph(b)
d.add_heading("Lista de cierre",1)
for x in ["Cuenta correcta verificada","Permisos mínimos revisados","Resultados contrastados con la fuente","Acciones sensibles confirmadas","Incidencias registradas","Acceso conservado o revocado con justificación"]:d.add_paragraph(x,style="List Bullet")
d.add_paragraph("Material educativo: no sustituye políticas internas, asesoría jurídica, controles de seguridad ni revisión profesional.")
d.save(OUT/"guia-seguridad-privacidad-revocacion.docx")

styles=getSampleStyleSheet(); T=ParagraphStyle("T",parent=styles["Title"],fontName="Helvetica-Bold",fontSize=19,textColor=colors.HexColor("#"+GOLD),alignment=1,spaceAfter=8); B=ParagraphStyle("B",parent=styles["BodyText"],fontName="Helvetica",fontSize=7.4,leading=9,textColor=colors.HexColor("#"+NAVY)); H=ParagraphStyle("H",parent=B,fontName="Helvetica-Bold",textColor=colors.white,alignment=1)
def pdf(name,title,headers,rows,widths,pagesize=letter):
 story=[Paragraph(title,T),Paragraph("ZYRON AI LAB · 28-07-2026",ParagraphStyle("S",parent=B,alignment=1)),Spacer(1,6)]
 data=[[Paragraph(str(x),H) for x in headers]]+[[Paragraph(str(x),B) for x in row] for row in rows]
 t=Table(data,colWidths=widths,repeatRows=1,splitByRow=1); t.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor("#"+BLUE)),("GRID",(0,0),(-1,-1),.35,colors.HexColor("#AAB8C5")),("VALIGN",(0,0),(-1,-1),"TOP"),("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white,colors.HexColor("#"+LIGHT)]),("LEFTPADDING",(0,0),(-1,-1),3),("RIGHTPADDING",(0,0),(-1,-1),3),("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3)])); story.append(t)
 SimpleDocTemplate(str(OUT/name),pagesize=pagesize,rightMargin=24,leftMargin=24,topMargin=28,bottomMargin=28).build(story)

allfields=[x for _,items in sections for x in items]
pdf("plantilla-planificacion-autorizacion-conectores.pdf","Plantilla de planificación y autorización de conectores",["Campo","Contenido"],[[x,""] for x in allfields],[170,350])
pdf("guia-seguridad-privacidad-revocacion.pdf","Guía de seguridad, privacidad y revocación",["Control","Aplicación"],security,[150,370])
matrix=["Gmail","Google Calendar","Google Drive","Google Contacts","Notion","Servicio empresarial compatible","Otra fuente autorizada"]
pdf("matriz-permisos-fuentes-acciones.pdf","Matriz de permisos, fuentes y acciones",["Fuente","Cuenta","Datos","Lectura","Escritura","Confirmación","Responsable","Revocación"],[[x,"","","","","","",""] for x in matrix],[90,70,110,65,65,75,75,80],landscape(letter))
rub=[("Definición del objetivo",10),("Selección de fuente",10),("Gestión de permisos",15),("Calidad de la consulta",15),("Verificación de resultados",15),("Seguridad y privacidad",15),("Registro y evidencias",10),("Validación y cierre",10)]
pdf("rubrica-laboratorio-conectores.pdf","Rúbrica del Laboratorio Conectores",["Criterio","Máx.","Excelente","Competente","En desarrollo","Requiere revisión","Obtenido"],[[a,b,"Dominio completo","Cumple lo esencial","Cumplimiento parcial","Debe repetirse",""] for a,b in rub],[110,35,110,110,110,110,55],landscape(letter))
headers=["ID","Fecha","Fuente","Cuenta","Consulta/acción","Permisos","Resultado","Confirmación","Incidencia","Corrección","Evidencia","Cierre"]
pdf("registro-consultas-acciones-incidencias-resultados.pdf","Registro de consultas, acciones, incidencias y resultados",headers,[[i+1,"","","","","","","","","","",""] for i in range(15)],[25,45,65,55,85,65,75,65,65,65,55,55],landscape(letter))
